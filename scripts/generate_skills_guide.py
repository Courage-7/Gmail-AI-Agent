"""Generate the project skills installation and usage guide."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DOCX = OUT_DIR / "email-agent-skills-guide.docx"

INK = RGBColor(24, 33, 47)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(71, 85, 105)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F8FAFC"
SUCCESS = "DCFCE7"
WARNING = "FEF3C7"


def main() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    configure_document(doc)

    add_title(doc)
    add_status_callout(doc)
    add_status_table(doc)
    add_section_overview(doc)
    add_installation_section(doc)
    add_start_section(doc)
    add_frontend_workflow_section(doc)
    add_operating_rules(doc)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("Email Agent Skills Guide")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = MUTED


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Email Agent Skills Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Installation status, project startup, and design workflow for the visual workflow builder")
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta_run = meta.add_run("Project: email-agent  |  Scope: Graphify, UI/UX Pro Max, Framer Motion")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = MUTED


def add_status_callout(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_BLUE)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.add_run("Current position: ").bold = True
    p.add_run(
        "Graphify and UI/UX Pro Max are installed in the user-level Codex skills folder. "
        "Framer Motion is planned for the frontend project, but the package download was blocked by the current approval limit."
    )


def add_status_table(doc: Document) -> None:
    doc.add_heading("1. Installation Status", level=1)
    table = doc.add_table(rows=1, cols=4)
    set_table_width(table, [1800, 1800, 2520, 3240])
    headers = ["Tool", "Status", "Where", "What it is for"]
    for idx, text in enumerate(headers):
        cell = table.cell(0, idx)
        shade_cell(cell, LIGHT_BLUE)
        set_cell_text(cell, text, bold=True)

    rows = [
        (
            "Graphify",
            "Installed",
            r"C:\Users\coura\.codex\skills\graphify",
            "Build and query a knowledge graph of the codebase before architecture-heavy work.",
        ),
        (
            "UI/UX Pro Max",
            "Installed",
            r"C:\Users\coura\.codex\skills\ui-ux-pro-max",
            "Design intelligence for professional frontend layout, hierarchy, colors, UX, and component decisions.",
        ),
        (
            "Framer Motion",
            "Pending package install",
            r"frontend/package.json",
            "React animation library for polished transitions, node states, panels, and workflow feedback.",
        ),
        (
            "21st.dev Magic MCP",
            "Skipped for now",
            "Needs API key",
            "Optional component-generation MCP for later; not installed in this pass.",
        ),
    ]
    for tool, status, where, purpose in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], tool, bold=True)
        set_cell_text(cells[1], status)
        shade_cell(cells[1], SUCCESS if status == "Installed" else WARNING)
        set_cell_text(cells[2], where)
        set_cell_text(cells[3], purpose)


def add_section_overview(doc: Document) -> None:
    doc.add_heading("2. What Each Tool Adds", level=1)
    add_bullet(doc, "Graphify turns the project into a queryable knowledge graph. Use it when we need codebase-level understanding before changing architecture, APIs, runtime flow, or component boundaries.")
    add_bullet(doc, "UI/UX Pro Max gives design guidance for frontend polish: information density, hierarchy, responsive behavior, professional styling, component choices, and UX heuristics.")
    add_bullet(doc, "Framer Motion will add production-ready animation primitives once installed in the Vite frontend. It should support usability, not decorate randomly.")
    add_bullet(doc, "21st.dev can be considered later if we want generated UI components from its MCP server and have an API key.")


def add_installation_section(doc: Document) -> None:
    doc.add_heading("3. Installation Guide", level=1)

    doc.add_heading("Graphify", level=2)
    add_code_block(
        doc,
        [
            "uv tool install graphifyy",
            r"C:\Users\coura\.local\bin\graphify.exe install --platform codex",
            r"C:\Users\coura\.local\bin\graphify.exe --version",
        ],
    )
    add_bullet(doc, r"Installed skill path: C:\Users\coura\.codex\skills\graphify")
    add_bullet(doc, r"Note: C:\Users\coura\.local\bin is not currently on PATH. Use the absolute graphify.exe path or update the shell PATH later.")

    doc.add_heading("UI/UX Pro Max", level=2)
    add_code_block(
        doc,
        [
            "npm install -g uipro-cli",
            "uipro init --ai codex",
            r"Copy project skill folder to C:\Users\coura\.codex\skills\ui-ux-pro-max if needed",
            "uipro --version",
        ],
    )
    add_bullet(doc, r"Installed skill path: C:\Users\coura\.codex\skills\ui-ux-pro-max")
    add_bullet(doc, r"The installer also created a project-local copy at .codex/skills/ui-ux-pro-max.")

    doc.add_heading("Framer Motion", level=2)
    add_code_block(
        doc,
        [
            "cd frontend",
            "npm install framer-motion",
            "npm run build",
        ],
    )
    add_bullet(doc, "Status: not installed yet because external package download approval is currently blocked.")
    add_bullet(doc, "Once installed, import from framer-motion in React components only where animation clarifies workflow state or interaction.")


def add_start_section(doc: Document) -> None:
    doc.add_heading("4. How To Start Them In This Project", level=1)

    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, [2100, 3600, 3660])
    for idx, text in enumerate(["Tool", "Start command or trigger", "Good first use in email-agent"]):
        cell = table.cell(0, idx)
        shade_cell(cell, LIGHT_BLUE)
        set_cell_text(cell, text, bold=True)

    rows = [
        (
            "Graphify",
            r"After Codex restart, invoke the graphify skill for codebase questions. CLI docs also show: /graphify .",
            "Map FastAPI routes, workflow registry, frontend components, and future LangGraph adapter boundaries.",
        ),
        (
            "UI/UX Pro Max",
            "After Codex restart, ask Codex to use ui-ux-pro-max for frontend design and UX polish.",
            "Audit the workflow builder layout, node palette, config panel, validation states, empty states, and responsive layout.",
        ),
        (
            "Framer Motion",
            "After package install: import { motion, AnimatePresence } from 'framer-motion'.",
            "Animate node status changes, validation feedback, panel transitions, and run-event output without slowing the tool UI.",
        ),
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text, bold=(idx == 0))


def add_frontend_workflow_section(doc: Document) -> None:
    doc.add_heading("5. Professional Frontend Workflow", level=1)
    doc.add_paragraph(
        "Yes, these tools can help us design the workflow-builder frontend professionally. The best order is:"
    )
    for item in [
        "Use Graphify when we need a map of the existing system before changing backend/frontend contracts.",
        "Use UI/UX Pro Max before a visual polish pass. It should guide layout density, panel hierarchy, typography, validation states, and responsive behavior.",
        "Install Framer Motion after the core builder and validation loop are stable. Add motion only where it improves feedback: drag/drop affordance, node run status, validation summary, and panel transitions.",
        "Keep React Flow as the interaction engine. Framer Motion should enhance surrounding UI and state transitions, not fight React Flow's canvas behavior.",
    ]:
        add_numbered(doc, item)

    doc.add_heading("Recommended animation rules", level=2)
    add_bullet(doc, "Prefer subtle 120-220 ms transitions for panels, alerts, and status badges.")
    add_bullet(doc, "Use clear state changes over decorative movement. The builder is an operational tool.")
    add_bullet(doc, "Respect reduced-motion settings.")
    add_bullet(doc, "Avoid animating every node constantly; reserve motion for run state, validation, and direct manipulation feedback.")


def add_operating_rules(doc: Document) -> None:
    doc.add_heading("6. Operating Rules Before We Continue", level=1)
    add_bullet(doc, "Restart Codex before expecting newly installed skills to appear in the available skills list.")
    add_bullet(doc, "Do not expose raw MCP, Gmail, or internal tool execution from the browser. The backend registry stays authoritative.")
    add_bullet(doc, "Use the skills as design and architecture accelerators, not as permission to skip tests, validation, or browser verification.")
    add_bullet(doc, "Next implementation step after restart: connect the builder to backend validation cleanly, then polish the UX with UI/UX Pro Max guidance.")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p.add_run(text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(lines):
        if index:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(15, 23, 42)


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.text = ""
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.color.rgb = INK


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


if __name__ == "__main__":
    main()
